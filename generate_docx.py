from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


def set_default_font(doc, font_name="Times New Roman", font_size=12):
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(6)


def add_heading_custom(doc, text, level=1):
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.name = "Times New Roman"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    if level == 0:
        run.font.size = Pt(20)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return heading


def add_paragraph(doc, text, bold=False, italic=False, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if align:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.15


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 1.15


def page_break(doc):
    doc.add_page_break()


# Create document
doc = Document()
set_default_font(doc)

# Title page
add_paragraph(doc, "ONLINE AUCTION SYSTEM", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Technical Documentation and Project Report", bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Final Year Project 2026", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Prepared by: [Student Name]", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Supervisor/Lecturer: [Lecturer Name]", align=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph(doc, "Date: 12 August 2026", align=WD_ALIGN_PARAGRAPH.CENTER)

page_break(doc)

# Table of Contents
add_heading_custom(doc, "Table of Contents", level=1)
add_paragraph(doc, "Chapter 1: Introduction ............................................................ 1")
add_paragraph(doc, "Chapter 2: Tools and Technologies ............................................... 4")
add_paragraph(doc, "Chapter 3: System Architecture and Database Design ............... 7")
add_paragraph(doc, "Chapter 4: System Features and Functionalities ...................... 10")
add_paragraph(doc, "Chapter 5: Innovations, Strengths and Future Recommendations . 14")

page_break(doc)

# ============================================
# CHAPTER 1
# ============================================
add_heading_custom(doc, "Chapter 1: Introduction", level=1)

add_heading_custom(doc, "1.1 Background", level=2)
add_paragraph(doc,
    "Online auction platforms have transformed the way second-hand goods, collectibles, "
    "art and specialist items are bought and sold. Unlike fixed-price marketplaces, auctions "
    "create competitive, time-bound environments where buyers determine the final price of an "
    "item through open bidding. This project implements a web-based auction system designed "
    "for collectors, dealers and everyday sellers who want a simple yet complete timed-auction "
    "experience.")

add_heading_custom(doc, "1.2 Problem Statement", level=2)
add_paragraph(doc,
    "Many small-scale auction solutions either lack proper role separation, provide no audit "
    "trail for administrative actions, or fail to automate the post-auction workflow such as "
    "invoice generation and payment tracking. This project addresses these gaps by building a "
    "role-aware platform with automatic lot closure, integrated invoicing, bid moderation and "
    "an admin audit log.")

add_heading_custom(doc, "1.3 Objectives", level=2)
add_bullets(doc, [
    "To develop a responsive web application for listing and bidding on timed auctions.",
    "To support three distinct user roles: buyers, sellers and administrators.",
    "To provide a secure authentication system with password hashing and session management.",
    "To automate the auction lifecycle including opening, bidding, closing and invoicing.",
    "To give administrators tools for lot, user, bid and payment management.",
    "To record an audit trail of administrative actions for accountability."
])

add_heading_custom(doc, "1.4 Scope and Delimitations", level=2)
add_paragraph(doc,
    "The system covers user registration and authentication, lot creation by sellers, public "
    "browsing and search, bidding, watchlists, automatic lot closure, invoice generation, payment "
    "recording, and comprehensive admin controls. It does not integrate with real payment gateways "
    "or email/SMS providers; payment settlement is recorded administratively. Real-time bidding "
    "via WebSockets is also outside the current scope.")

add_heading_custom(doc, "1.5 Target Users", level=2)
add_bullets(doc, [
    "Buyers: registered users who browse lots, place bids, manage watchlists and pay invoices.",
    "Sellers: registered users who consign items, upload photographs and track bidding activity.",
    "Administrators: users who moderate lots, bids and accounts, and settle payments on behalf of buyers."
])

page_break(doc)

# ============================================
# CHAPTER 2
# ============================================
add_heading_custom(doc, "Chapter 2: Tools and Technologies", level=1)

add_heading_custom(doc, "2.1 Overview of the Technology Stack", level=2)
add_paragraph(doc,
    "The project is a server-rendered web application that follows a Model-View-Controller "
    "pattern. Python and Flask form the server-side layer, Jinja2 templates render the views, "
    "MySQL stores persistent data, and Bootstrap together with custom CSS and JavaScript handle "
    "the user interface.")

add_heading_custom(doc, "2.2 Backend Technologies", level=2)
add_bullets(doc, [
    "Python 3.12: the core programming language used for the application logic.",
    "Flask 3.1: a lightweight web framework that handles routing, sessions and request processing.",
    "Flask-Login 0.6: manages user session state and the current_user proxy.",
    "Flask-MySQLdb 2.0 and mysqlclient 2.2: connect the Flask application to the MySQL database.",
    "Werkzeug 3.1: provides password hashing, request utilities and the development server."
])

add_heading_custom(doc, "2.3 Frontend Technologies", level=2)
add_bullets(doc, [
    "HTML5 and Jinja2: server-rendered markup with reusable templates and macros.",
    "Bootstrap 5.3: responsive layout grid, forms, modals and navigation components.",
    "Custom CSS: a design-system layer with colour tokens, typography and components.",
    "JavaScript (vanilla): client-side countdowns, charts, form validation, image previews and navigation.",
    "Font Awesome 6.4: iconography used throughout the interface.",
    "Chart.js 4.4: renders dashboard area charts for bids, revenue and activity.",
    "Google Fonts (Inter, Plus Jakarta Sans): body and display typefaces."
])

add_heading_custom(doc, "2.4 Database Management System", level=2)
add_paragraph(doc,
    "The persistent store is a MySQL or MariaDB relational database named auction_db, using the "
    "utf8mb4 character set. The schema is created and migrated automatically when the application "
    "starts. Separate tables hold users, auction lots, lot images, bids, payments, watchlist "
    "entries and an admin audit log.")

add_heading_custom(doc, "2.5 Development and Testing Environment", level=2)
add_bullets(doc, [
    "Operating system: Microsoft Windows.",
    "Code editor: Visual Studio Code with a Chrome launch configuration.",
    "Virtual environment: an isolated Python venv folder keeps dependencies separate from the system.",
    "Database administration: phpMyAdmin or the MySQL command-line client.",
    "Utility scripts: image seeder, image optimiser, smoke tests and sanity checks."
])

page_break(doc)

# ============================================
# CHAPTER 3
# ============================================
add_heading_custom(doc, "Chapter 3: System Architecture and Database Design", level=1)

add_heading_custom(doc, "3.1 Architectural Pattern", level=2)
add_paragraph(doc,
    "The system follows a lightweight Model-View-Controller architecture. Routes act as controllers, "
    "accepting HTTP requests, applying business rules and returning rendered Jinja2 views. The model "
    "layer is implemented through SQL queries against the MySQL database rather than an Object "
    "Relational Mapper, which keeps the data access explicit and easy to follow.")

add_heading_custom(doc, "3.2 System Components and Modules", level=2)
add_paragraph(doc,
    "The main application file contains all route handlers, helper functions and business logic. "
    "A separate database operations module is responsible for creating the database, tables and "
    "default administrator account when the system first runs. Templates are organised into a base "
    "public layout, a dashboard workspace layout, page-specific templates and shared macros for "
    "lot cards, charts and statistics. Static assets include the stylesheet, JavaScript file and "
    "uploaded lot images or user avatars.")

add_heading_custom(doc, "3.3 Database Schema", level=2)
add_paragraph(doc,
    "The relational schema is designed to support the complete auction lifecycle. The seven tables "
    "and their purposes are summarised below.")

# Table of tables
table = doc.add_table(rows=1, cols=2)
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Table"
hdr_cells[1].text = "Purpose"
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.bold = True
            run.font.size = Pt(12)

rows = [
    ("users", "Stores account details, roles, profile images and suspension status."),
    ("auctions", "Stores lot details, pricing, schedule, status, bid count and winner."),
    ("auction_images", "Links photographs to lots and marks the primary catalogue image."),
    ("bids", "Records each bid, its status, and moderation history including removal reason."),
    ("payments", "Represents invoices or payments, their status and settlement details."),
    ("watchlist", "Tracks saved lots for each user."),
    ("admin_actions", "Audit log of administrative write operations.")
]

for label, purpose in rows:
    row_cells = table.add_row().cells
    row_cells[0].text = label
    row_cells[1].text = purpose
    for cell in row_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

add_heading_custom(doc, "3.4 Data Flow", level=2)
add_paragraph(doc,
    "A seller creates a lot by entering a title, department, description, opening bid, minimum "
    "increment and closing time. The lot becomes active and appears on the public home page, "
    "catalogue and search. Buyers view the lot, place bids that must satisfy the minimum "
    "increment, and may add the lot to a watchlist. When the closing time passes, the system "
    "automatically selects the highest accepted bid, declares a winner and generates a pending "
    "payment. The buyer can then settle the invoice, or an administrator can record payment on "
    "their behalf. All admin actions are written to the audit log.")

page_break(doc)

# ============================================
# CHAPTER 4
# ============================================
add_heading_custom(doc, "Chapter 4: System Features and Functionalities", level=1)

add_heading_custom(doc, "4.1 Public Features", level=2)
add_bullets(doc, [
    "Home page: a landing view with live lot count, featured lot and department links.",
    "Showcase landing page: a full-screen, horizontally scrolling presentation of every lot.",
    "Browse page: a department catalogue with pagination, price bands, status filters and sorting.",
    "Search page: keyword, category, price, status, bid count, ending time and seller filters.",
    "Lot detail page: full description, image gallery, bid history, seller details and countdown timer."
])

add_heading_custom(doc, "4.2 Authentication and User Roles", level=2)
add_paragraph(doc,
    "Users register with a username, email, password, first name, last name, phone and a chosen "
    "role of buyer or seller. Passwords are hashed before storage. During login, the role is "
    "checked and the user is directed to the appropriate dashboard. Access to seller, buyer and "
    "admin routes is protected by role checks.")

add_heading_custom(doc, "4.3 Buyer Features", level=2)
add_bullets(doc, [
    "Place bids on active lots with automatic validation of minimum increments.",
    "Receive feedback when already the highest bidder or when a duplicate bid is submitted.",
    "Maintain a watchlist of interesting lots.",
    "View a personal dashboard showing bids, watchlist, wins, pending invoices and spending history.",
    "Pay invoices via the payment form.",
    "View and print invoice pages with buyer, seller and lot details."
])

add_heading_custom(doc, "4.4 Seller Features", level=2)
add_bullets(doc, [
    "Consign new lots with title, category, description, opening bid, bid increment and closing time.",
    "Upload multiple photographs; the first image is used as the primary catalogue photo.",
    "Track all consigned lots, live lots, closed lots, watchers and total revenue.",
    "View dashboard charts showing bidding activity and earnings over the previous fourteen days."
])

add_heading_custom(doc, "4.5 Administrative Features", level=2)
add_bullets(doc, [
    "Overview dashboard with key performance indicators, charts and attention alerts.",
    "List, search, edit, activate, close, cancel or permanently delete lots.",
    "Search, suspend, unsuspend and change the role of any user.",
    "Review all bids, remove suspicious or invalid bids with a recorded reason, and restore them if needed.",
    "Raise invoices for closed lots, record payments and reopen payments if necessary.",
    "View an audit log of all admin actions with timestamps."
])

add_heading_custom(doc, "4.6 Profile and Account Management", level=2)
add_bullets(doc, [
    "Edit first name, last name, phone and profile image.",
    "View personal activity statistics including bids placed, lots consigned, auctions won and watchlist size.",
    "Close own account with confirmation; the system blocks closure when invoices are unsettled, "
    "active lots have bids, or the account is the only administrator."
])

page_break(doc)

# ============================================
# CHAPTER 5
# ============================================
add_heading_custom(doc, "Chapter 5: Innovations, Strengths and Future Recommendations", level=1)

add_heading_custom(doc, "5.1 Key Innovations", level=2)
add_numbered(doc, [
    "Automatic auction closure and invoice generation: the system evaluates expired active lots, "
    "selects the highest accepted bid and immediately creates a pending payment for the winner.",
    "Bid moderation with transparent reasoning: administrators can remove or restore bids. Removed "
    "bids carry a reason that the bidder can see, and the lot price and winner are recalculated "
    "automatically.",
    "Comprehensive audit logging: every administrative write is recorded, including the admin, "
    "action, target and note.",
    "User suspension and role management: administrators can suspend accounts and promote or "
    "demote users between buyer, seller and admin roles.",
    "Safe self-service account deletion: the system protects active auctions and unpaid invoices "
    "and prevents deletion of the only administrator.",
    "Immersive showcase landing page: a full-screen, keyboard-navigable horizontal presentation "
    "of every lot with thumbnail galleries and countdown timers.",
    "Custom, responsive design system: a tailored visual layer over Bootstrap provides consistent "
    "typography, colour tokens, micro-interactions and accessibility considerations.",
    "Dashboard analytics: fourteen-day area charts of bidding activity, revenue and spending are "
    "rendered using Chart.js.",
    "Automated image seeding and optimisation: demo lots can be populated with openly licensed "
    "images from the Openverse API, then resized and re-encoded to keep page loads fast.",
    "Outbid tracking and live statistics: buyer and seller dashboards show not only totals but "
    "also the number of lots a buyer is currently losing.",
    "Graceful client-side degradation: JavaScript initialisers are isolated so a single failed "
    "widget cannot break the whole page."
])

add_heading_custom(doc, "5.2 System Strengths", level=2)
add_bullets(doc, [
    "Clear separation of roles and responsibilities between buyers, sellers and administrators.",
    "Server-rendered pages keep the application simple and accessible without requiring a separate front-end build step.",
    "The use of raw SQL makes the data flow transparent and easy to tune.",
    "Built-in utilities for seeding data, optimising images and smoke-testing the application.",
    "Consistent, modern user interface with countdowns, counters, charts and responsive layouts."
])

add_heading_custom(doc, "5.3 Limitations and Security Considerations", level=2)
add_bulples = [
    "The application runs with a hard-coded secret key and plaintext database credentials in the source code; these must be moved to environment variables for production.",
    "Payment settlement is recorded manually or administratively; no real payment gateway is integrated.",
    "There is no email or SMS notification system for outbids, auction closure or payment reminders.",
    "Bidding is request-based rather than real-time; simultaneous bids are serialised by the database but there is no live WebSocket feed."
]
add_bullets(doc, add_bulples)

add_heading_custom(doc, "5.4 Future Enhancements", level=2)
add_numbered(doc, [
    "Externalise configuration: store the secret key, database credentials and other settings in environment variables.",
    "Integrate a real payment gateway such as PayPal, Stripe or a local mobile-money provider for automatic settlement.",
    "Add email notifications for registration, outbid alerts, auction closure and invoice reminders.",
    "Implement WebSocket or server-sent events for real-time bid updates and countdowns.",
    "Introduce automated testing with unit tests and integration tests for the route layer.",
    "Add reporting and export features for administrators, including PDF summaries and CSV downloads.",
    "Support image CDN or cloud storage to reduce local disk usage and improve delivery speed."
])

add_heading_custom(doc, "5.5 Conclusion", level=2)
add_paragraph(doc,
    "This project demonstrates a complete, role-based online auction platform built with widely "
    "used open-source tools. It goes beyond a basic bidding site by automating the auction lifecycle, "
    "providing bid moderation, audit logging, rich dashboards and a polished, responsive user "
    "interface. With the recommended security and integration improvements, the system could serve "
    "as a foundation for a production auction service.")

# Save
doc_path = r"C:\Users\mmoyana\OneDrive - Champions Insurance Zimbabwe\Desktop\final project 2026 modified\final project 2026\Auction_System_Documentation.docx"
doc.save(doc_path)
print(f"Saved: {doc_path}")
